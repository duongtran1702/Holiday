from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== STYLE SETUP =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
    return table

def add_data_dict_table(doc, table_name, fields):
    """fields: list of (field, type, desc)"""
    p = doc.add_paragraph()
    run = p.add_run(f'Bảng: {table_name}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    add_table(doc, ['Trường', 'Kiểu dữ liệu', 'Mô tả'], fields)
    doc.add_paragraph('')

# ===== TITLE PAGE =====
for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = 'Times New Roman'

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub1.add_run('(Software Requirements Specification - SRS)')
run.font.size = Pt(14)
run.italic = True
run.font.name = 'Times New Roman'

doc.add_paragraph('')

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('HỆ THỐNG QUẢN LÝ VÀ KINH DOANH THỜI TRANG TÍCH HỢP ABC')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = 'Times New Roman'

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub3.add_run('Bán lẻ trực tuyến (B2C) và Quản lý đại lý/khách sỉ phân phối (B2B)\ncho Cửa hàng thời trang ABC')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')

info_items = [
    ('Phiên bản:', '1.0'),
    ('Ngày phát hành:', '08/07/2026'),
    ('Trạng thái:', 'Dự thảo (Draft)'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p.add_run(label + ' ')
    run_l.bold = True
    run_l.font.size = Pt(13)
    run_l.font.name = 'Times New Roman'
    run_v = p.add_run(value)
    run_v.font.size = Pt(13)
    run_v.font.name = 'Times New Roman'



# ===== 1. GIỚI THIỆU =====
doc.add_heading('1. GIỚI THIỆU', level=1)

doc.add_heading('1.1 Mục đích tài liệu', level=2)
doc.add_paragraph(
    'Tài liệu này đặc tả đầy đủ các yêu cầu nghiệp vụ và chức năng của Hệ thống Thương mại điện tử B2C/B2B '
    'phục vụ hoạt động bán lẻ quần áo trực tuyến cho khách hàng cá nhân, quản lý đại lý/khách sỉ phân phối '
    'và kiểm soát tồn kho sản phẩm thời trang tại cửa hàng ABC. '
    'Tài liệu là cơ sở để đội ngũ phát triển thiết kế, xây dựng, kiểm thử hệ thống, đồng thời là căn cứ '
    'để các bên liên quan (chủ cửa hàng, đội vận hành, đội kỹ thuật) thống nhất phạm vi và tiêu chí nghiệm thu.'
)

doc.add_heading('1.2 Phạm vi hệ thống', level=2)
doc.add_paragraph('Hệ thống bao gồm 3 nhóm phân hệ chính:')
doc.add_paragraph(
    'Phân hệ B2C (Cổng Bán lẻ): phục vụ khách hàng cá nhân (như anh Mynato) mua sắm quần áo trực tuyến với giá niêm yết, '
    'thanh toán tiền tươi qua cổng thanh toán hoặc COD.',
    style='List Bullet'
)
doc.add_paragraph(
    'Phân hệ B2B (Cổng Bán sỉ): phục vụ đại lý, chủ buôn, đối tác kinh doanh (như ông ATMIN) đặt hàng số lượng lớn '
    'theo cơ chế chiết khấu bậc thang và chính sách công nợ.',
    style='List Bullet'
)
doc.add_paragraph(
    'Phân hệ Quản trị (Admin Dashboard): quản lý sản phẩm/biến thể, tồn kho, đơn hàng cả hai luồng, '
    'đại lý, khuyến mãi và báo cáo doanh thu/dòng tiền.',
    style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run('Hệ thống không bao gồm: ')
run.bold = True
run.font.name = 'Times New Roman'
p.add_run(
    'hạ tầng mạng/máy chủ vật lý, hệ thống POS bán hàng tại quầy, hệ thống quản trị nội bộ khác '
    '(HRM, Accounting tổng, BI toàn doanh nghiệp), hệ thống logistics/vận chuyển nội bộ '
    '— các hệ thống này được xem là ngoài phạm vi (out of scope) hoặc chỉ tích hợp ở mức API.'
).font.name = 'Times New Roman'

doc.add_heading('1.3 Đối tượng sử dụng tài liệu', level=2)
doc.add_paragraph('Chủ cửa hàng / Ban quản lý dự án: đối chiếu phạm vi, phê duyệt yêu cầu.', style='List Bullet')
doc.add_paragraph('Đội phát triển phần mềm (Full-stack Developer): làm căn cứ thiết kế và lập trình.', style='List Bullet')
doc.add_paragraph('Đội kiểm thử (QA/Tester): xây dựng test case dựa trên yêu cầu.', style='List Bullet')
doc.add_paragraph('Đơn vị vận hành cửa hàng: xác nhận nghiệp vụ đúng với thực tế vận hành.', style='List Bullet')

doc.add_heading('1.4 Định nghĩa, từ viết tắt', level=2)
add_table(doc,
    ['Thuật ngữ', 'Giải thích'],
    [
        ['B2C', 'Business to Customer – Giao dịch giữa cửa hàng và khách hàng cá nhân mua lẻ.'],
        ['B2B', 'Business to Business – Giao dịch giữa cửa hàng và đại lý/khách sỉ.'],
        ['Agent (Đại lý/Khách sỉ)', 'Đối tác phân phối hoặc chủ buôn, được cấp tài khoản B2B và mức chiết khấu riêng.'],
        ['SRS', 'Software Requirements Specification – Tài liệu đặc tả yêu cầu phần mềm.'],
        ['SKU', 'Stock Keeping Unit – Đơn vị quản lý tồn kho, ứng với một biến thể cụ thể (sản phẩm + size + màu).'],
        ['RBAC', 'Role-Based Access Control – Kiểm soát truy cập theo vai trò.'],
        ['Matrix Order', 'Biểu mẫu đặt hàng dạng ma trận (Size × Màu) để nhập số lượng nhanh cho nhiều biến thể.'],
        ['Công nợ đại lý', 'Số tiền đại lý/khách sỉ còn nợ hệ thống sau khi đặt hàng theo hình thức ghi nợ.'],
        ['Chiết khấu bậc thang', 'Chính sách giá giảm dần theo số lượng mua: mua càng nhiều, giá càng rẻ.'],
        ['Tồn kho (Inventory)', 'Số lượng sản phẩm còn lại có thể bán theo từng SKU (Size × Màu).'],
        ['Voucher', 'Mã giảm giá áp dụng khi mua hàng trên cổng B2C.'],
        ['COD', 'Cash on Delivery – Thanh toán khi nhận hàng.'],
    ]
)

doc.add_heading('1.5 Tài liệu tham khảo', level=2)
doc.add_paragraph('Bảng giá và chính sách chiết khấu bậc thang nội bộ của cửa hàng ABC.', style='List Bullet')
doc.add_paragraph('Tài liệu API cổng thanh toán VNPay / Momo Sandbox.', style='List Bullet')
doc.add_paragraph('Quy chuẩn UI/UX nội bộ (nếu có).', style='List Bullet')
doc.add_paragraph('Danh mục sản phẩm và cấu trúc phân loại quần áo hiện hành.', style='List Bullet')

# ===== 2. MÔ TẢ TỔNG QUAN HỆ THỐNG =====

doc.add_heading('2. MÔ TẢ TỔNG QUAN HỆ THỐNG', level=1)

doc.add_heading('2.1 Bối cảnh nghiệp vụ', level=2)
doc.add_paragraph(
    'Cửa hàng thời trang ABC hiện vận hành bán hàng theo hình thức truyền thống (tại cửa hàng vật lý) '
    'kết hợp một phần trực tuyến qua các kênh phân phối riêng lẻ (fanpage Facebook, Zalo, Shopee), '
    'chưa có hệ thống thống nhất để: '
    '(1) khách hàng cá nhân tự xem hàng, đặt mua và thanh toán trực tuyến trên website riêng; '
    '(2) đại lý/khách sỉ đặt hàng số lượng lớn với chính sách giá chiết khấu bậc thang và quản lý công nợ tự động; '
    '(3) chủ shop theo dõi doanh thu, tồn kho theo từng biến thể (size × màu) và công nợ đại lý theo thời gian thực.'
)
doc.add_paragraph(
    'Hệ thống B2C/B2B được xây dựng nhằm số hóa toàn bộ quy trình bán hàng và phân phối sỉ, '
    'giảm phụ thuộc vào thao tác thủ công (nhắn tin báo giá, ghi sổ tay công nợ), '
    'tăng khả năng kiểm soát tồn kho và mở rộng kênh bán hàng chuyên nghiệp.'
)

doc.add_heading('2.2 Sơ đồ tổng quan chức năng hệ thống', level=2)
doc.add_paragraph(
    'Hệ thống được tổ chức theo mô hình 3 lớp chức năng chính, tương tác với nhau qua một cơ sở dữ liệu tập trung:'
)
doc.add_paragraph('Lớp giao dịch B2C – Website bán lẻ quần áo cho khách hàng cá nhân.', style='List Number')
doc.add_paragraph('Lớp giao dịch B2B – Cổng đặt hàng sỉ dành cho đại lý/chủ buôn.', style='List Number')
doc.add_paragraph('Lớp quản trị nội bộ – Admin Dashboard (Chủ shop, Nhân viên kho, Kế toán, Sales).', style='List Number')

doc.add_paragraph('')
# ASCII architecture diagram
arch = '''                          +------------------------------------------+
                          |        HỆ THỐNG SHOP THỜI TRANG ABC       |
                          +------------------------------------------+
                                              |
               +------------------------------+------------------------------+
               |                              |                              |
   [CỔNG BÁN LẺ - B2C]            [CỔNG BÁN SỈ - B2B]          [ADMIN DASHBOARD]
 - Khách: Anh Mynato             - Khách: Ông ATMIN            - Chủ shop ABC
 - Mua lẻ: 1-5 bộ               - Mua sỉ: >= 10 bộ           - Nhân viên kho
 - Giá: Niêm yết cố định        - Giá: Chiết khấu bậc thang  - Kế toán
 - TT: Momo/VNPAY/COD           - TT: Ký quỹ, Ghi sổ nợ     - Sales
 - UX: Đẹp, Voucher, Review     - UX: Matrix Order Form       - Báo cáo BI, Phân quyền
 - Live Chat Hỗ trợ Khách hàng  - Live Chat Hỗ trợ Khách hàng - Trả lời Live Chat'''
p = doc.add_paragraph(arch)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
doc.add_paragraph('')

doc.add_paragraph(
    'Cả 3 lớp cùng thao tác trên một tập dữ liệu dùng chung: sản phẩm/biến thể, tồn kho theo SKU, '
    'đơn hàng, thanh toán, công nợ, người dùng — đảm bảo tính nhất quán dữ liệu theo thời gian thực (real-time) '
    'giữa các kênh bán.'
)

doc.add_heading('2.3 Các nhóm người dùng (Actors)', level=2)
add_table(doc,
    ['Vai trò', 'Mô tả', 'Quyền hạn chính'],
    [
        ['Khách hàng lẻ\n(B2C Customer)',
         'Khách hàng cá nhân, vãng lai hoặc hội viên, mua quần áo số lượng nhỏ.',
         'Xem sản phẩm, thêm giỏ hàng, đặt hàng, thanh toán, đánh giá sản phẩm, xem lịch sử đơn hàng'],
        ['Đại lý / Khách sỉ\n(B2B Partner)',
         'Chủ shop, đại lý, tổ chức doanh nghiệp đã được Admin phê duyệt.',
         'Đặt hàng số lượng lớn theo giá chiết khấu bậc thang, sử dụng Matrix Order Form, ghi nợ, xem công nợ, xem báo cáo doanh số'],
        ['Nhân viên Sales',
         'Nhân viên vận hành nội bộ của cửa hàng ABC.',
         'Xử lý đơn hàng, hỗ trợ khách hàng qua Live Chat, quản lý khuyến mãi'],
        ['Nhân viên Kho',
         'Nhân viên quản lý kho hàng vật lý.',
         'Xem đơn hàng, cập nhật tồn kho, không có quyền xóa dữ liệu'],
        ['Kế toán',
         'Nhân viên kế toán nội bộ.',
         'Quản lý công nợ, xem báo cáo, không có quyền can thiệp sản phẩm/đơn hàng'],
        ['Quản trị viên\n(Admin)',
         'Chủ cửa hàng ABC hoặc quản lý cấp cao.',
         'Toàn quyền: quản lý sản phẩm, người dùng, đại lý, khuyến mãi, xem báo cáo BI, phân quyền'],
    ]
)

doc.add_heading('2.4 Giả định và ràng buộc', level=2)
doc.add_paragraph(
    'Hệ thống sử dụng cổng thanh toán VNPay/Momo (môi trường Sandbox) trong giai đoạn thử nghiệm, '
    'có thể thay thế bằng cổng chính thức khi vận hành thật.',
    style='List Bullet'
)
doc.add_paragraph(
    'Chính sách chiết khấu bậc thang được áp dụng thống nhất cho tất cả đại lý trong giai đoạn 1 '
    '(chưa phân hạng đại lý Vàng/Bạc/Đồng); việc phân hạng là hướng mở rộng.',
    style='List Bullet'
)
doc.add_paragraph(
    'Mỗi sản phẩm quần áo có nhiều biến thể (variant) theo tổ hợp Size × Màu sắc. '
    'Tồn kho được quản lý tới cấp SKU (một biến thể = một SKU).',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống giả định người dùng có kết nối Internet ổn định khi thực hiện giao dịch.',
    style='List Bullet'
)
doc.add_paragraph(
    'Vận chuyển đơn hàng được xử lý ngoài hệ thống (qua đơn vị vận chuyển bên thứ 3); '
    'hệ thống chỉ cập nhật trạng thái đơn hàng.',
    style='List Bullet'
)
doc.add_paragraph(
    'Ngôn ngữ hệ thống: Tiếng Việt (giai đoạn 1); có thể mở rộng đa ngôn ngữ.',
    style='List Bullet'
)

# ===== 3. YÊU CẦU NGHIỆP VỤ CHI TIẾT =====

doc.add_heading('3. YÊU CẦU NGHIỆP VỤ CHI TIẾT', level=1)

# --- 3.1 B2C ---
doc.add_heading('3.1 Nhóm chức năng B2C (Khách hàng cá nhân)', level=2)

doc.add_heading('3.1.1 Đăng ký / Đăng nhập', level=3)
doc.add_paragraph(
    'Mô tả: Khách hàng có thể đăng ký tài khoản, đăng nhập, hoặc mua hàng không cần tài khoản (Guest checkout).'
)
doc.add_paragraph('Đăng ký bằng email + mật khẩu, xác thực qua email/OTP.', style='List Bullet')
doc.add_paragraph('Đăng nhập bằng email/mật khẩu; hỗ trợ đăng nhập nhanh qua Google (tùy chọn mở rộng).', style='List Bullet')
doc.add_paragraph(
    'Cho phép đặt hàng không cần đăng nhập (Guest), yêu cầu nhập họ tên, email, số điện thoại, địa chỉ giao hàng.',
    style='List Bullet'
)
doc.add_paragraph('Quên mật khẩu: gửi liên kết đặt lại mật khẩu qua email.', style='List Bullet')

doc.add_heading('3.1.2 Tra cứu, xem sản phẩm quần áo', level=3)
doc.add_paragraph('Hiển thị danh sách sản phẩm quần áo theo danh mục (Áo, Quần, Đầm/Váy, Phụ kiện, ...).', style='List Bullet')
doc.add_paragraph(
    'Hiển thị thông tin chi tiết sản phẩm: hình ảnh (nhiều góc), giá niêm yết, mô tả chất liệu, '
    'hướng dẫn chọn size, các biến thể có sẵn (Size × Màu).',
    style='List Bullet'
)
doc.add_paragraph('Hiển thị tình trạng còn hàng/hết hàng theo từng biến thể (SKU).', style='List Bullet')
doc.add_paragraph('Tìm kiếm và lọc sản phẩm theo: danh mục, khoảng giá, size, màu sắc, từ khóa.', style='List Bullet')
doc.add_paragraph('Sắp xếp sản phẩm theo: mới nhất, giá tăng/giảm, bán chạy nhất.', style='List Bullet')

doc.add_heading('3.1.3 Giỏ hàng', level=3)
doc.add_paragraph('Thêm sản phẩm vào giỏ hàng với biến thể cụ thể (Size + Màu) và số lượng.', style='List Bullet')
doc.add_paragraph('Xóa hoặc cập nhật số lượng sản phẩm trong giỏ hàng.', style='List Bullet')
doc.add_paragraph('Hiển thị tổng tạm tính, cho phép nhập mã giảm giá (Voucher) nếu có.', style='List Bullet')
doc.add_paragraph('Kiểm tra tồn kho theo SKU thời gian thực trước khi cho phép thanh toán.', style='List Bullet')
doc.add_paragraph(
    'Hệ thống khóa kho tạm thời (Lock Inventory) cho các SKU trong giỏ hàng khi khách bắt đầu thanh toán '
    'để tránh tranh chấp dữ liệu (Concurrency Control).',
    style='List Bullet'
)

doc.add_heading('3.1.4 Đặt hàng & Thanh toán', level=3)
doc.add_paragraph(
    'Quy trình: Chọn sản phẩm (Size/Màu) → Giỏ hàng → Nhập thông tin giao hàng → '
    'Chọn phương thức thanh toán → Thanh toán → Xác nhận đơn hàng.'
)
doc.add_paragraph(
    'Phương thức thanh toán hỗ trợ: Thanh toán trực tuyến qua VNPay/Momo, Thanh toán khi nhận hàng (COD).',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống tạo đơn hàng ở trạng thái "Chờ thanh toán" trước khi chuyển hướng sang cổng thanh toán.',
    style='List Bullet'
)
doc.add_paragraph(
    'Sau khi nhận callback/webhook từ cổng thanh toán, hệ thống cập nhật trạng thái đơn hàng: '
    'Đã thanh toán / Thanh toán thất bại.',
    style='List Bullet'
)
doc.add_paragraph(
    'Nếu thanh toán thất bại hoặc đơn hàng "Chờ thanh toán" quá 15 phút, hệ thống tự động hủy đơn '
    'và giải phóng tồn kho (release inventory) để bán cho khách khác.',
    style='List Bullet'
)
doc.add_paragraph(
    'Luồng trừ kho và luồng tạo đơn hàng phải nằm trong cùng một Database Transaction. '
    'Nếu một trong hai bước lỗi, hệ thống phải Rollback toàn bộ.',
    style='List Bullet'
)

doc.add_heading('3.1.5 Đánh giá sản phẩm', level=3)
doc.add_paragraph(
    'Khách hàng đã mua và nhận hàng thành công có quyền để lại đánh giá cho sản phẩm đã mua.',
    style='List Bullet'
)
doc.add_paragraph('Đánh giá bao gồm: số sao (1-5), nội dung bình luận, ảnh thực tế (tùy chọn).', style='List Bullet')
doc.add_paragraph('Hiển thị đánh giá trung bình và danh sách review trên trang chi tiết sản phẩm.', style='List Bullet')

doc.add_heading('3.1.6 Quản lý đơn hàng cá nhân', level=3)
doc.add_paragraph(
    'Xem lịch sử đơn hàng đã đặt, trạng thái đơn hàng: Chờ thanh toán → Đã thanh toán → '
    'Đang xử lý → Đang giao → Đã giao → Hoàn thành.',
    style='List Bullet'
)
doc.add_paragraph('Xem chi tiết từng đơn hàng: danh sách sản phẩm, số lượng, giá, tổng tiền, thông tin giao hàng.', style='List Bullet')
doc.add_paragraph('Hủy đơn hàng (chỉ khi đơn ở trạng thái "Chờ xử lý", chưa đóng gói).', style='List Bullet')

# --- 3.2 B2B ---

doc.add_heading('3.2 Nhóm chức năng B2B (Đại lý / Khách sỉ)', level=2)

doc.add_heading('3.2.1 Đăng ký & Xét duyệt tài khoản doanh nghiệp', level=3)
doc.add_paragraph(
    'Đối tác/khách sỉ đăng ký tài khoản đại lý qua biểu mẫu riêng biệt: thông tin doanh nghiệp/cửa hàng, '
    'mã số thuế (nếu có), tên người đại diện, số điện thoại, email, địa chỉ.',
    style='List Bullet'
)
doc.add_paragraph(
    'Tài khoản đại lý ở trạng thái "Chờ duyệt" cho đến khi Admin/Sales xác nhận. '
    'Đại lý không thể đăng nhập giao dịch khi chưa được duyệt.',
    style='List Bullet'
)
doc.add_paragraph(
    'Sau khi duyệt, hệ thống cấp: tài khoản đăng nhập cổng B2B, hạn mức công nợ (credit limit).',
    style='List Bullet'
)
doc.add_paragraph(
    'Admin có thể tạm khóa / ngừng hợp tác với đại lý bất kỳ lúc nào.',
    style='List Bullet'
)

doc.add_heading('3.2.2 Chính sách giá sỉ – Chiết khấu bậc thang', level=3)
doc.add_paragraph(
    'Hệ thống tự động tính giá dựa trên tổng số lượng sản phẩm trong đơn hàng, áp dụng chính sách chiết khấu bậc thang:'
)
p = doc.add_paragraph()
run = p.add_run('Ví dụ minh họa (có thể cấu hình bởi Admin):')
run.italic = True
run.font.name = 'Times New Roman'

add_table(doc,
    ['Bậc', 'Số lượng mua', 'Đơn giá', 'Ghi chú'],
    [
        ['Bậc 1', '1 – 9 bộ', '200.000 đ/bộ', 'Giá niêm yết (không chiết khấu)'],
        ['Bậc 2', '10 – 49 bộ', '150.000 đ/bộ', 'Giảm 25%'],
        ['Bậc 3', '50 – 99 bộ', '120.000 đ/bộ', 'Giảm 40%'],
        ['Bậc 4', '≥ 100 bộ', '100.000 đ/bộ', 'Giảm 50%'],
    ]
)
doc.add_paragraph('')
doc.add_paragraph(
    'Giá hiển thị cho đại lý khi đặt hàng là giá đã áp dụng chiết khấu bậc thang, '
    'tự động cập nhật khi đại lý thay đổi số lượng. Khác hoàn toàn với giá hiển thị cho khách B2C.',
    style='List Bullet'
)
doc.add_paragraph(
    'Admin có thể tạo nhiều bảng chiết khấu bậc thang khác nhau cho từng danh mục sản phẩm hoặc theo mùa.',
    style='List Bullet'
)

doc.add_heading('3.2.3 Đặt hàng theo Ma trận (Matrix Order Form)', level=3)
doc.add_paragraph(
    'Đây là tính năng cốt lõi dành riêng cho B2B, cho phép đại lý nhập nhanh số lượng đặt hàng '
    'của tất cả các biến thể (Size × Màu) trên một bảng duy nhất, thay vì phải thêm từng biến thể vào giỏ hàng.'
)
p = doc.add_paragraph()
run = p.add_run('Ví dụ minh họa Matrix Order Form cho sản phẩm "Áo Polo ABC Classic":')
run.italic = True
run.font.name = 'Times New Roman'

add_table(doc,
    ['Size \\ Màu', 'Trắng', 'Đen', 'Xanh Navy', 'Xám', 'Tổng/Size'],
    [
        ['S', '5', '10', '0', '5', '20'],
        ['M', '10', '15', '10', '10', '45'],
        ['L', '10', '20', '15', '5', '50'],
        ['XL', '5', '10', '5', '0', '20'],
        ['Tổng/Màu', '30', '55', '30', '20', '135 bộ'],
    ]
)
doc.add_paragraph('')
doc.add_paragraph(
    'Đại lý chỉ cần nhập số vào ô tương ứng, hệ thống tự tính tổng và áp dụng giá chiết khấu bậc thang '
    'dựa trên tổng số lượng toàn đơn (ví dụ: 135 bộ → Bậc 4: 100.000 đ/bộ).',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống kiểm tra tồn kho theo từng SKU (Size × Màu) trước khi xác nhận đặt hàng. '
    'Nếu SKU nào hết hàng, hệ thống cảnh báo ngay trên ô tương ứng.',
    style='List Bullet'
)

doc.add_heading('3.2.4 Thanh toán & Ghi nợ đại lý', level=3)
doc.add_paragraph(
    'Đại lý có thể chọn một trong hai hình thức thanh toán cho mỗi đơn hàng:',
)
doc.add_paragraph(
    '(a) Thanh toán ngay: Toàn bộ giá trị đơn hàng được thanh toán trực tuyến qua cổng thanh toán hoặc chuyển khoản.',
    style='List Bullet'
)
doc.add_paragraph(
    '(b) Ghi nợ (Credit): Đặt hàng trước, thanh toán sau. Hệ thống ghi nhận vào sổ công nợ. '
    'Đại lý có thể trả trước một phần (ví dụ 30%), phần còn lại ghi nợ.',
    style='List Bullet'
)
doc.add_paragraph(
    'Đơn hàng ghi nợ chỉ được xác nhận nếu: Tổng công nợ hiện tại + Giá trị đơn hàng mới ≤ Hạn mức công nợ được cấp.',
    style='List Bullet'
)

doc.add_heading('3.2.5 Quản lý Công nợ', level=3)
doc.add_paragraph(
    'Hệ thống ghi nhận từng đơn hàng ghi nợ vào bảng công nợ đại lý, kèm hạn thanh toán.',
    style='List Bullet'
)
doc.add_paragraph(
    'Đại lý xem được: tổng công nợ hiện tại, hạn mức còn lại, danh sách đơn hàng chưa thanh toán, lịch sử thanh toán.',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống tự động gửi email nhắc nợ định kỳ khi công nợ gần đến hạn hoặc quá hạn.',
    style='List Bullet'
)
doc.add_paragraph(
    'Admin/Kế toán xác nhận thanh toán công nợ (cập nhật thủ công hoặc đối soát chuyển khoản), '
    'cập nhật trạng thái "Đã thanh toán".',
    style='List Bullet'
)

doc.add_heading('3.2.6 Báo cáo doanh số đại lý', level=3)
doc.add_paragraph(
    'Đại lý xem được báo cáo: số đơn hàng đã đặt, tổng số lượng sản phẩm, doanh số theo tháng, '
    'tổng chiết khấu được hưởng.',
    style='List Bullet'
)

# --- 3.3 ADMIN ---

doc.add_heading('3.3 Nhóm chức năng Quản trị (Admin Dashboard)', level=2)

doc.add_heading('3.3.1 Quản lý sản phẩm & Biến thể (Catalog/Inventory)', level=3)
doc.add_paragraph(
    'Thêm/sửa/xóa/ẩn sản phẩm quần áo. Mỗi sản phẩm bao gồm: tên, mô tả, chất liệu, hình ảnh, '
    'danh mục, giá niêm yết, trạng thái (hiển thị/ẩn).',
    style='List Bullet'
)
doc.add_paragraph(
    'Quản lý biến thể (Variant): Mỗi sản phẩm có nhiều biến thể theo tổ hợp Size × Màu sắc. '
    'Mỗi biến thể tương ứng một SKU riêng.',
    style='List Bullet'
)
doc.add_paragraph(
    'Quản lý tồn kho theo từng SKU (Size × Màu). Cập nhật số lượng tồn kho khi nhập hàng mới.',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống khóa kho tạm thời (Lock Inventory) cho các SKU đang trong quá trình thanh toán '
    'để tránh tranh chấp dữ liệu (Concurrency Control) giữa nhiều khách mua cùng lúc.',
    style='List Bullet'
)
doc.add_paragraph(
    'Theo dõi số lượng tồn kho theo thời gian thực, cảnh báo khi sản phẩm sắp hết hàng.',
    style='List Bullet'
)

doc.add_heading('3.3.2 Quản lý đơn hàng (Order Fulfillment)', level=3)
doc.add_paragraph(
    'Xem, tìm kiếm, lọc toàn bộ đơn hàng B2C và B2B theo trạng thái, thời gian, kênh bán, khách hàng.',
    style='List Bullet'
)
doc.add_paragraph(
    'Cập nhật trạng thái đơn hàng theo quy trình: Chờ duyệt → Đang đóng gói → '
    'Đã giao cho đơn vị vận chuyển → Đã giao → Hoàn thành.',
    style='List Bullet'
)
doc.add_paragraph(
    'Xử lý yêu cầu hủy/hoàn tiền đơn hàng (nếu chính sách cho phép, trước khi đóng gói).',
    style='List Bullet'
)
doc.add_paragraph(
    'Xem chi tiết từng đơn hàng: sản phẩm + biến thể, khách hàng/đại lý, trạng thái thanh toán.',
    style='List Bullet'
)

doc.add_heading('3.3.3 Quản lý khuyến mãi (Voucher/Promotion)', level=3)
doc.add_paragraph(
    'Tạo mã giảm giá (Voucher): theo % hoặc số tiền cố định, giới hạn thời gian áp dụng, '
    'giới hạn số lượt sử dụng, giá trị đơn hàng tối thiểu.',
    style='List Bullet'
)
doc.add_paragraph(
    'Áp dụng khuyến mãi cho toàn bộ hoặc một số danh mục/sản phẩm nhất định.',
    style='List Bullet'
)
doc.add_paragraph(
    'Chỉ áp dụng cho kênh B2C (đại lý B2B đã có chính sách chiết khấu bậc thang riêng).',
    style='List Bullet'
)

doc.add_heading('3.3.4 Quản lý đại lý', level=3)
doc.add_paragraph('Duyệt/từ chối đăng ký đại lý mới, kèm ghi chú lý do (nếu từ chối).', style='List Bullet')
doc.add_paragraph('Thiết lập/điều chỉnh hạn mức công nợ cho từng đại lý.', style='List Bullet')
doc.add_paragraph('Cấu hình bảng chiết khấu bậc thang áp dụng cho kênh B2B.', style='List Bullet')
doc.add_paragraph('Khóa/mở khóa tài khoản đại lý.', style='List Bullet')

doc.add_heading('3.3.5 Quản lý công nợ đại lý (dành cho Kế toán)', level=3)
doc.add_paragraph('Xem tổng hợp công nợ toàn bộ đại lý: tổng nợ, nợ quá hạn, nợ trong hạn.', style='List Bullet')
doc.add_paragraph('Xem chi tiết công nợ từng đại lý: danh sách đơn hàng ghi nợ, hạn thanh toán, trạng thái.', style='List Bullet')
doc.add_paragraph('Xác nhận thanh toán công nợ khi nhận được chuyển khoản, cập nhật trạng thái "Đã thanh toán".', style='List Bullet')

doc.add_heading('3.3.6 Báo cáo & Thống kê (BI Dashboard)', level=3)
doc.add_paragraph('Báo cáo doanh thu theo ngày/tuần/tháng, phân tách theo kênh B2C và B2B.', style='List Bullet')
doc.add_paragraph('Báo cáo top sản phẩm bán chạy, sản phẩm tồn kho lâu.', style='List Bullet')
doc.add_paragraph('Báo cáo dự báo mặt hàng sắp hết trong kho (Low Stock Alert).', style='List Bullet')
doc.add_paragraph('Biểu đồ trực quan (dashboard) tổng hợp các chỉ số kinh doanh.', style='List Bullet')
doc.add_paragraph('Báo cáo so sánh doanh thu B2B vs B2C theo thời gian.', style='List Bullet')

doc.add_heading('3.3.7 Quản trị người dùng & Ma trận Phân quyền (RBAC Matrix)', level=3)
doc.add_paragraph(
    'Hệ thống cung cấp Ma trận Phân quyền (Permission Matrix) với các quyền chi tiết (View, Create, Update, Delete) '
    'cho từng phân hệ (Products, Inventory, Orders, Agents, Debts, Promotions, Reports, Inbox).',
    style='List Bullet'
)
doc.add_paragraph(
    'Hỗ trợ cấu hình nhanh cho các Role: Nhân viên Kho (chỉ quản lý kho, đơn hàng), Sales (quản lý đơn hàng, đại lý, khuyến mãi, hỗ trợ khách), '
    'Kế toán (quản lý công nợ, xem báo cáo).',
    style='List Bullet'
)
doc.add_paragraph(
    'Admin toàn quyền điều chỉnh chi tiết từng quyền hạn của từng nhân viên. '
    'Giao diện chặn quyền truy cập (Forbidden 403) hiển thị khi truy cập trái phép.',
    style='List Bullet'
)

doc.add_heading('3.3.8 Hỗ trợ khách hàng (Live Chat / Inbox)', level=3)
doc.add_paragraph(
    'Tích hợp Widget Chat trực tuyến tại cổng B2C và B2B cho phép khách hàng trò chuyện trực tiếp với nhân viên.',
    style='List Bullet'
)
doc.add_paragraph(
    'Admin Inbox: Bảng điều khiển riêng dành cho nhân viên/admin nhận và phản hồi tin nhắn của khách hàng theo thời gian thực.',
    style='List Bullet'
)

# ===== 4. QUY TẮC NGHIỆP VỤ =====

doc.add_heading('4. QUY TẮC NGHIỆP VỤ (BUSINESS RULES)', level=1)
add_table(doc,
    ['Mã', 'Quy tắc', 'Áp dụng cho'],
    [
        ['BR-01',
         'Đơn hàng "Chờ thanh toán" quá 15 phút sẽ tự động hủy và giải phóng tồn kho.',
         'B2C, B2B'],
        ['BR-02',
         'Luồng trừ kho và tạo đơn hàng phải nằm trong cùng một Database Transaction. '
         'Nếu một bước lỗi, hệ thống phải Rollback toàn bộ.',
         'B2C, B2B'],
        ['BR-03',
         'Đại lý chỉ được đặt hàng ghi nợ nếu (Công nợ hiện tại + Giá trị đơn hàng) ≤ Hạn mức công nợ được cấp.',
         'B2B'],
        ['BR-04',
         'Giá bán cho đại lý được tính theo bảng chiết khấu bậc thang, dựa trên tổng số lượng trong đơn hàng.',
         'B2B'],
        ['BR-05',
         'Không cho phép đặt hàng vượt quá số lượng tồn kho còn lại của SKU tương ứng (Size × Màu).',
         'B2C, B2B'],
        ['BR-06',
         'Mã giảm giá (Voucher) chỉ áp dụng khi còn trong thời hạn hiệu lực, chưa vượt số lượt sử dụng tối đa, '
         'và chỉ áp dụng cho kênh B2C.',
         'B2C'],
        ['BR-07',
         'Tài khoản đại lý mới phải ở trạng thái "Chờ duyệt" và không thể đăng nhập giao dịch cho đến khi được Admin duyệt.',
         'B2B'],
        ['BR-08',
         'Khách hàng chỉ được đánh giá sản phẩm đã mua và nhận hàng thành công.',
         'B2C'],
        ['BR-09',
         'Đơn hàng chỉ được hủy khi ở trạng thái "Chờ xử lý" (chưa đóng gói).',
         'B2C, B2B'],
        ['BR-10',
         'Hệ thống tự động gửi email nhắc nợ khi công nợ đại lý gần đến hạn hoặc quá hạn thanh toán.',
         'B2B'],
    ]
)

# ===== 5. YÊU CẦU PHI CHỨC NĂNG =====

doc.add_heading('5. YÊU CẦU PHI CHỨC NĂNG', level=1)

doc.add_heading('5.1 Hiệu năng (Performance)', level=2)
doc.add_paragraph(
    'Thời gian tải trang chủ B2C và trang danh sách sản phẩm không quá 2 giây trong điều kiện mạng bình thường, '
    'để tránh khách hàng rời bỏ trang web.',
    style='List Bullet'
)
doc.add_paragraph(
    'Sử dụng Cache (như Redis) cho các dữ liệu ít thay đổi: danh mục sản phẩm, cấu hình trang web, bảng chiết khấu — '
    'để giảm tải cho Database.',
    style='List Bullet'
)
doc.add_paragraph(
    'Hệ thống xử lý đồng thời tối thiểu 100 giao dịch đặt hàng/phút ở giai đoạn cao điểm '
    '(mục tiêu thiết kế, không bắt buộc kiểm thử tải đầy đủ trong giai đoạn MVP).',
    style='List Bullet'
)

doc.add_heading('5.2 Bảo mật (Security)', level=2)
doc.add_paragraph(
    'Mật khẩu người dùng được mã hóa (hashing) trước khi lưu trữ (sử dụng bcrypt hoặc tương đương).',
    style='List Bullet'
)
doc.add_paragraph(
    'Phân quyền truy cập API theo vai trò (RBAC): '
    'đảm bảo khách lẻ B2C không thể truy cập cổng xem giá sỉ B2B, '
    'nhân viên kho không thể can thiệp module kế toán/công nợ.',
    style='List Bullet'
)
doc.add_paragraph(
    'Giao dịch thanh toán tuân thủ cơ chế callback/webhook có xác thực chữ ký (checksum) từ cổng thanh toán.',
    style='List Bullet'
)
doc.add_paragraph('Toàn bộ kết nối sử dụng HTTPS.', style='List Bullet')

doc.add_heading('5.3 Tính toàn vẹn dữ liệu (Data Consistency)', level=2)
doc.add_paragraph(
    'Luồng trừ kho và tạo đơn hàng phải nằm trong cùng một Database Transaction. '
    'Nếu một trong hai bước lỗi, hệ thống phải Rollback lại toàn bộ '
    '(tránh: khách bị trừ tiền nhưng kho không trừ, đơn không tạo).',
    style='List Bullet'
)
doc.add_paragraph(
    'Dữ liệu tồn kho phải nhất quán giữa các kênh bán (B2C và B2B) khi có giao dịch phát sinh.',
    style='List Bullet'
)

doc.add_heading('5.4 Khả năng mở rộng (Scalability)', level=2)
doc.add_paragraph(
    'Kiến trúc cho phép bổ sung thêm kênh bán (Mobile App, Kiosk tại cửa hàng) thông qua API dùng chung.',
    style='List Bullet'
)
doc.add_paragraph(
    'Thiết kế dữ liệu cho phép mở rộng đa cấp đại lý trong tương lai mà không cần thay đổi cấu trúc lớn.',
    style='List Bullet'
)

doc.add_heading('5.5 Khả năng sử dụng (Usability)', level=2)
doc.add_paragraph(
    'Giao diện B2C responsive, tối ưu cho thiết bị di động (mobile-first), trải nghiệm mua sắm thân thiện.',
    style='List Bullet'
)
doc.add_paragraph(
    'Quy trình đặt hàng B2C không quá 4 bước thao tác chính.',
    style='List Bullet'
)
doc.add_paragraph(
    'Giao diện B2B (Matrix Order Form) tối ưu cho việc nhập liệu nhanh số lượng lớn trên desktop.',
    style='List Bullet'
)

doc.add_heading('5.6 Độ tin cậy & Sẵn sàng (Reliability)', level=2)
doc.add_paragraph(
    'Dữ liệu giao dịch (đơn hàng, thanh toán, công nợ) phải được lưu trữ nhất quán, '
    'không thất thoát khi xảy ra lỗi trong quá trình thanh toán.',
    style='List Bullet'
)
doc.add_paragraph(
    'Không có đơn hàng "treo" không xác định trạng thái.',
    style='List Bullet'
)

# ===== 6. MÔ HÌNH DỮ LIỆU =====

doc.add_heading('6. MÔ HÌNH DỮ LIỆU (DATA REQUIREMENTS)', level=1)

doc.add_heading('6.1 Danh sách thực thể chính', level=2)
add_table(doc,
    ['Thực thể', 'Mô tả'],
    [
        ['users', 'Tài khoản người dùng (khách hàng B2C, nhân viên, admin), lưu trữ quyền hạn (permissions) dưới dạng JSON.'],
        ['chat_messages', 'Lưu trữ nội dung tin nhắn giữa khách hàng và bộ phận hỗ trợ (Live Chat).'],
        ['agents', 'Thông tin đại lý/khách sỉ: mức chiết khấu, hạn mức công nợ, trạng thái duyệt.'],
        ['categories', 'Danh mục sản phẩm (Áo, Quần, Đầm/Váy, Phụ kiện...).'],
        ['products', 'Sản phẩm quần áo: tên, mô tả, chất liệu, hình ảnh, giá niêm yết.'],
        ['product_variants', 'Biến thể sản phẩm theo tổ hợp Size × Màu sắc, mỗi biến thể là một SKU.'],
        ['inventory', 'Tồn kho theo từng SKU (product_variant_id), bao gồm số lượng available và locked.'],
        ['promotions', 'Mã giảm giá (Voucher) và điều kiện áp dụng cho kênh B2C.'],
        ['pricing_tiers', 'Bảng chiết khấu bậc thang cho kênh B2B (bậc, khoảng số lượng, giá/chiết khấu).'],
        ['orders', 'Đơn hàng (B2C và B2B), trạng thái đơn hàng, trạng thái thanh toán.'],
        ['order_items', 'Chi tiết từng sản phẩm/biến thể trong đơn hàng, số lượng, đơn giá.'],
        ['payments', 'Giao dịch thanh toán liên kết với đơn hàng.'],
        ['agent_debts', 'Công nợ phát sinh từ đơn hàng ghi nợ của đại lý.'],
        ['reviews', 'Đánh giá sản phẩm của khách hàng B2C (sao, bình luận, hình ảnh).'],
    ]
)

doc.add_heading('6.2 Từ điển dữ liệu (Data Dictionary) — các bảng chính', level=2)

add_data_dict_table(doc, 'products', [
    ['id', 'UUID', 'Khóa chính'],
    ['category_id', 'UUID (FK)', 'Liên kết đến danh mục sản phẩm'],
    ['name', 'String', 'Tên sản phẩm'],
    ['description', 'Text', 'Mô tả chi tiết sản phẩm'],
    ['material', 'String', 'Chất liệu vải'],
    ['base_price', 'Decimal', 'Giá niêm yết (giá bán lẻ B2C)'],
    ['images', 'JSON/Array', 'Danh sách đường dẫn hình ảnh sản phẩm'],
    ['status', 'Enum', 'Active / Hidden / Discontinued'],
    ['created_at', 'Datetime', 'Thời điểm tạo'],
])

add_data_dict_table(doc, 'product_variants', [
    ['id', 'UUID', 'Khóa chính (= SKU ID)'],
    ['product_id', 'UUID (FK)', 'Liên kết đến sản phẩm gốc'],
    ['size', 'String', 'Size: S, M, L, XL, XXL...'],
    ['color', 'String', 'Màu sắc: Trắng, Đen, Xanh Navy...'],
    ['sku_code', 'String', 'Mã SKU duy nhất (ví dụ: ABC-POLO-WH-M)'],
    ['stock_quantity', 'Integer', 'Số lượng tồn kho khả dụng'],
    ['locked_quantity', 'Integer', 'Số lượng đang bị khóa tạm (đang trong quá trình thanh toán)'],
])

add_data_dict_table(doc, 'orders', [
    ['id', 'UUID', 'Khóa chính'],
    ['user_id / agent_id', 'UUID (FK)', 'Người mua (khách B2C) hoặc đại lý B2B'],
    ['order_type', 'Enum', 'B2C hoặc B2B'],
    ['status', 'Enum', 'Chờ thanh toán / Đã thanh toán / Đang xử lý / Đang đóng gói / Đã giao ĐVVC / Đã giao / Hoàn thành / Đã hủy'],
    ['payment_method', 'Enum', 'VNPay / Momo / COD / Ghi nợ (chỉ B2B)'],
    ['subtotal', 'Decimal', 'Tổng giá trị trước giảm giá'],
    ['discount_amount', 'Decimal', 'Số tiền giảm giá (Voucher hoặc chiết khấu bậc thang)'],
    ['total_amount', 'Decimal', 'Tổng giá trị đơn hàng sau giảm giá'],
    ['shipping_address', 'Text', 'Địa chỉ giao hàng'],
    ['note', 'Text', 'Ghi chú đơn hàng'],
    ['created_at', 'Datetime', 'Thời điểm tạo đơn'],
])

add_data_dict_table(doc, 'agent_debts', [
    ['id', 'UUID', 'Khóa chính'],
    ['agent_id', 'UUID (FK)', 'Đại lý liên quan'],
    ['order_id', 'UUID (FK)', 'Đơn hàng ghi nợ tương ứng'],
    ['amount', 'Decimal', 'Số tiền công nợ'],
    ['paid_amount', 'Decimal', 'Số tiền đã trả trước (nếu có)'],
    ['due_date', 'Date', 'Hạn thanh toán'],
    ['paid_status', 'Enum', 'Chưa thanh toán / Thanh toán một phần / Đã thanh toán'],
    ['paid_at', 'Datetime', 'Thời điểm thanh toán hoàn tất (nếu có)'],
])

add_data_dict_table(doc, 'chat_messages', [
    ['id', 'UUID', 'Khóa chính'],
    ['sender_id', 'UUID (FK)', 'Người gửi (Khách hoặc Nhân viên)'],
    ['receiver_id', 'UUID (FK)', 'Người nhận (Admin hoặc Khách)'],
    ['content', 'Text', 'Nội dung tin nhắn'],
    ['timestamp', 'Datetime', 'Thời gian gửi'],
])

add_data_dict_table(doc, 'reviews', [
    ['id', 'UUID', 'Khóa chính'],
    ['user_id', 'UUID (FK)', 'Khách hàng đánh giá'],
    ['product_id', 'UUID (FK)', 'Sản phẩm được đánh giá'],
    ['order_id', 'UUID (FK)', 'Đơn hàng liên quan (chứng minh đã mua)'],
    ['rating', 'Integer', 'Số sao (1-5)'],
    ['comment', 'Text', 'Nội dung bình luận'],
    ['images', 'JSON/Array', 'Ảnh thực tế kèm theo (tùy chọn)'],
    ['created_at', 'Datetime', 'Thời điểm đánh giá'],
])

# ===== 7. USE CASE =====
doc.add_heading('7. DANH SÁCH USE CASE TỔNG HỢP', level=1)
add_table(doc,
    ['Mã', 'Tên Use Case', 'Actor', 'Mô tả ngắn'],
    [
        ['UC-01', 'Đăng ký / Đăng nhập', 'Khách hàng B2C', 'Tạo tài khoản hoặc đăng nhập hệ thống'],
        ['UC-02', 'Xem danh sách & chi tiết sản phẩm', 'Khách hàng B2C', 'Tra cứu, lọc, xem thông tin quần áo'],
        ['UC-03', 'Thêm giỏ hàng & Đặt hàng', 'Khách hàng B2C', 'Chọn sản phẩm (Size/Màu), thanh toán'],
        ['UC-04', 'Thanh toán trực tuyến', 'Khách hàng B2C', 'Thanh toán qua VNPay/Momo hoặc COD'],
        ['UC-05', 'Đánh giá sản phẩm', 'Khách hàng B2C', 'Để lại đánh giá sao, bình luận, ảnh'],
        ['UC-06', 'Áp dụng Voucher giảm giá', 'Khách hàng B2C', 'Nhập mã giảm giá khi thanh toán'],
        ['UC-07', 'Đăng ký tài khoản đại lý', 'Đại lý B2B', 'Gửi hồ sơ đăng ký đối tác phân phối'],
        ['UC-08', 'Đặt hàng sỉ (Matrix Order)', 'Đại lý B2B', 'Nhập nhanh số lượng theo bảng Size × Màu'],
        ['UC-09', 'Thanh toán / Ghi nợ đại lý', 'Đại lý B2B', 'Thanh toán ngay hoặc ghi nợ trong hạn mức'],
        ['UC-10', 'Xem công nợ đại lý', 'Đại lý B2B', 'Theo dõi công nợ, hạn mức còn lại'],
        ['UC-11', 'Duyệt đại lý mới', 'Admin / Sales', 'Xét duyệt hồ sơ đăng ký đại lý'],
        ['UC-12', 'Quản lý sản phẩm & tồn kho', 'Admin', 'CRUD sản phẩm, biến thể, tồn kho theo SKU'],
        ['UC-13', 'Quản lý đơn hàng', 'Admin / Sales / Kho', 'Xem, xử lý đơn hàng B2C/B2B'],
        ['UC-14', 'Quản lý khuyến mãi', 'Admin', 'Tạo và quản lý mã giảm giá Voucher'],
        ['UC-15', 'Xác nhận thanh toán công nợ', 'Kế toán', 'Xác nhận đại lý đã thanh toán nợ'],
        ['UC-16', 'Xem báo cáo BI', 'Admin', 'Xem dashboard doanh thu B2C/B2B, tồn kho'],
        ['UC-17', 'Phân quyền người dùng', 'Admin', 'Quản lý tài khoản nội bộ, phân quyền RBAC'],
    ]
)

# ===== 8. YÊU CẦU GIAO DIỆN =====
doc.add_heading('8. YÊU CẦU GIAO DIỆN (TÓM TẮT)', level=1)
doc.add_paragraph(
    'Giao diện B2C (Cổng Bán lẻ): Thiết kế thân thiện, hiện đại, responsive (mobile-first); '
    'ưu tiên hình ảnh sản phẩm chất lượng cao, quy trình đặt hàng tối giản, '
    'tích hợp mã giảm giá và hiển thị đánh giá sản phẩm. '
    'Trải nghiệm mua sắm phải đẹp, chuyên nghiệp, tạo cảm giác tin cậy.',
    style='List Bullet'
)
doc.add_paragraph(
    'Giao diện B2B (Cổng Đại lý): Tập trung vào chức năng nghiệp vụ, '
    'Matrix Order Form nhập liệu nhanh dạng bảng Size × Màu, '
    'hiển thị giá chiết khấu bậc thang rõ ràng, dashboard công nợ và doanh số.',
    style='List Bullet'
)
doc.add_paragraph(
    'Giao diện Admin Dashboard: Dạng dashboard quản trị, có biểu đồ trực quan (doanh thu, tồn kho), '
    'bảng dữ liệu có tìm kiếm/lọc/phân trang, giao diện quản lý sản phẩm + biến thể trực quan.',
    style='List Bullet'
)

# ===== 9. TIÊU CHÍ NGHIỆM THU =====
doc.add_heading('9. TIÊU CHÍ NGHIỆM THU TỔNG QUÁT', level=1)
criteria = [
    'Khách hàng B2C có thể hoàn tất một giao dịch đặt hàng từ đầu đến cuối '
    '(chọn sản phẩm + Size/Màu → giỏ hàng → thanh toán → nhận xác nhận) mà không gặp lỗi.',
    'Đại lý đã được duyệt có thể đặt hàng sỉ qua Matrix Order Form, '
    'hệ thống áp dụng đúng giá chiết khấu bậc thang và kiểm tra đúng hạn mức công nợ.',
    'Admin có thể xem báo cáo doanh thu tổng hợp, phân tách đúng theo kênh B2C/B2B.',
    'Tồn kho sản phẩm (theo SKU: Size × Màu) được cập nhật chính xác, nhất quán giữa các kênh bán '
    'khi có giao dịch phát sinh.',
    'Hệ thống khóa kho tạm thời (Lock Inventory) hoạt động đúng, tránh tranh chấp dữ liệu '
    'khi nhiều khách mua cùng lúc.',
    'Luồng trừ kho + tạo đơn hàng nằm trong cùng một Transaction; '
    'nếu lỗi giữa chừng, hệ thống Rollback toàn bộ.',
    'Toàn bộ giao dịch thanh toán được ghi nhận đầy đủ trạng thái, '
    'không có đơn hàng "treo" không xác định trạng thái.',
    'Phân quyền RBAC hoạt động đúng: khách B2C không thấy giá sỉ, '
    'nhân viên kho không truy cập được module công nợ.',
]
for i, c in enumerate(criteria, 1):
    doc.add_paragraph(f'{i}. {c}')

# ===== 10. PHỤ LỤC =====
doc.add_heading('10. PHỤ LỤC', level=1)

doc.add_heading('10.1 Kiến nghị công nghệ (Tech Stack)', level=2)
doc.add_paragraph(
    'Để triển khai bản SRS này một cách chuẩn chuyên môn, dưới đây là Stack công nghệ được đề xuất:'
)
doc.add_paragraph(
    'Backend: Java Spring Boot (Spring Security để phân quyền RBAC, Spring Data JPA để xử lý dữ liệu).',
    style='List Bullet'
)
doc.add_paragraph(
    'Database: MySQL hoặc PostgreSQL (xử lý các mối quan hệ thực thể chặt chẽ) kết hợp với Redis làm Cache.',
    style='List Bullet'
)
doc.add_paragraph(
    'Frontend: ReactJS hoặc Next.js — Dùng TailwindCSS để build giao diện responsive cho B2C, '
    'và dùng các thư viện UI như Ant Design/Shadcn để làm bảng biểu Dashboard phức tạp cho B2B/Admin.',
    style='List Bullet'
)

doc.add_heading('10.2 Hướng phát triển mở rộng (ngoài phạm vi giai đoạn 1)', level=2)
doc.add_paragraph('Phân hạng đại lý nhiều cấp (Vàng/Bạc/Đồng) với chính sách chiết khấu khác nhau.', style='List Bullet')
doc.add_paragraph('Chương trình tích điểm thành viên (Loyalty) cho khách hàng B2C.', style='List Bullet')
doc.add_paragraph('Tích hợp CMS quản lý nội dung trang chủ, bài viết, lookbook thời trang.', style='List Bullet')
doc.add_paragraph('Xuất báo cáo đối soát công nợ định dạng Excel/PDF tự động.', style='List Bullet')
doc.add_paragraph('Tích hợp hệ thống vận chuyển (GHN, GHTK, Viettel Post) để tracking đơn hàng.', style='List Bullet')
doc.add_paragraph('Ứng dụng di động (Mobile App) bán hàng B2C.', style='List Bullet')
doc.add_paragraph('Hệ thống POS (Point of Sale) bán hàng tại cửa hàng vật lý, đồng bộ tồn kho.', style='List Bullet')

doc.add_heading('10.3 Lịch sử thay đổi tài liệu', level=2)
add_table(doc,
    ['Phiên bản', 'Ngày', 'Người thực hiện', 'Nội dung thay đổi'],
    [
        ['1.0', '08/07/2026', 'Nhóm phát triển',
         'Khởi tạo tài liệu SRS đầy đủ nghiệp vụ B2C/B2B cho hệ thống thời trang ABC'],
    ]
)

# ===== SAVE =====
doc.save('D:\\Holiday\\SRS.docx')
print('SRS.docx saved successfully!')
